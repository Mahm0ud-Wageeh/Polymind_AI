from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\Projects\Graduation project\Polymind-fullstack")
OUTPUT = ROOT / "deliverables" / "Polymind_Graduation_Project_Proposal.docx"


def set_run_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def add_body_paragraph(document, label, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    return paragraph


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True)
    return paragraph


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("Graduation Project Proposal Form")
    set_run_font(title_run, size=14, bold=True)

    add_section_heading(document, "1. Project Information")
    add_body_paragraph(
        document,
        "Project Title: ",
        "Polymind - An AI-Powered Workspace for Network Engineering",
    )
    add_body_paragraph(document, "Course/Track: ", "Computer Science - Software Engineering Track")
    add_body_paragraph(document, "Team Members: ", "")
    for _ in range(7):
        member = document.add_paragraph(style="List Paragraph")
        member.paragraph_format.space_after = Pt(2)
        run = member.add_run("[Student Name] - ")
        set_run_font(run)

    add_section_heading(document, "2. Project Overview")
    add_body_paragraph(
        document,
        "Objective: ",
        "To develop a full-stack SaaS workspace that helps network engineers use artificial intelligence to design, document, analyze, troubleshoot, and manage IT infrastructure. The platform provides authenticated workspaces, projects, AI conversations, and reusable engineering resources in one environment.",
    )
    add_body_paragraph(
        document,
        "Scope of Work: ",
        "The project includes a React and TypeScript web application, a Laravel REST and streaming API, PostgreSQL data storage, Redis-backed cache and queues, workspace and project management, secure user authentication, role-based access control, AI chat with file attachments, conversation history, an engineering dashboard, and a library for agents and prompt templates. The first release focuses on network engineering workflows while preserving a modular structure for future cloud, DevOps, and cybersecurity modules. Native mobile applications and hardware-device management are outside the scope of this release.",
    )
    add_body_paragraph(
        document,
        "Expected Outcomes: ",
        "A working Docker-deployable platform where users can register, sign in, create and manage workspaces and projects, upload supporting files, ask an AI assistant for network-engineering help through streamed responses, keep persistent conversation history, and view real usage information. The final system will include automated tests for key frontend and backend flows.",
    )

    add_section_heading(document, "3. Problem Statement")
    add_body_paragraph(
        document,
        "",
        "Network engineers routinely switch between topology notes, device configuration references, troubleshooting records, documentation tools, and general-purpose AI chats. This fragments important context, makes collaboration and traceability difficult, and slows down repetitive engineering tasks. General AI tools also do not provide a dedicated workspace for organizing engineering projects, attached files, conversation history, access roles, and usage controls. There is a need for a secure, integrated platform that brings these activities together while allowing engineers to use reliable AI providers through one consistent interface.",
    )

    add_section_heading(document, "4. Proposed Solution")
    add_body_paragraph(
        document,
        "Technologies Used: ",
        "Frontend: React 18, TypeScript, Vite, Tailwind CSS, Zustand, and Radix UI. Backend: PHP 8.3 with Laravel 12. Database: PostgreSQL with UUID-based records. Cache, queues, and sessions: Redis. Authentication and authorization: Laravel Sanctum, Socialite OAuth, and Spatie Permission. AI: a provider-agnostic layer supporting OpenAI-compatible providers, Anthropic, Gemini, Groq, OpenRouter, DeepSeek, Mistral, and Ollama, with retry and fallback handling. Real-time AI replies use Server-Sent Events. Development and deployment use Docker Compose, Nginx, and GitHub Actions.",
    )
    add_body_paragraph(
        document,
        "System Architecture: ",
        "The React client communicates with a versioned Laravel API. Sanctum protects user sessions and API requests, while the backend provisions organizations and workspaces, enforces ownership and role permissions, and stores users, projects, conversations, messages, uploaded files, usage records, agents, and templates in PostgreSQL. Redis supports caching, queued tasks, sessions, and rate limiting. The AI manager selects the configured provider, sends requests, streams assistant responses to the client when supported, records token usage and cost information, and applies retries or provider fallbacks when a request fails. Docker containers run the frontend, API, PostgreSQL, and Redis as a reproducible deployment stack.",
    )

    add_section_heading(document, "5. Resources Needed")
    add_body_paragraph(
        document,
        "Hardware/Software: ",
        "A development computer with Docker Desktop, internet access for configured AI providers, and optional cloud hosting for deployment. Required software includes PHP 8.3 and Composer, Node.js and npm, Docker Compose, PostgreSQL, Redis, Git and GitHub, and a modern web browser. The project may use Figma for interface planning and an AI provider account for live model access during evaluation.",
    )

    add_section_heading(document, "6. Approval")
    add_body_paragraph(document, "Instructor/Advisor: ", "............................................................")
    add_body_paragraph(document, "Signature: ", "........................................................................")

    document.core_properties.title = "Polymind Graduation Project Proposal"
    document.core_properties.subject = "Graduation project proposal"
    document.core_properties.author = "Polymind Project Team"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
